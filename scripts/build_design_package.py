from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUTPUT = DOCS / "Horse_Asset_Manager_設計資料一式.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "1C2733"
MUTED = "5F6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "B58120"
RED = "9B1C1C"
GREEN = "2E6B4E"

ASCII_FONT = "Calibri"
EAST_ASIA_FONT = "Yu Gothic"
MONO_FONT = "Consolas"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    ascii_name: str = ASCII_FONT,
    east_asia: str = EAST_ASIA_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = ascii_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_name)
    r_fonts.set(qn("w:hAnsi"), ascii_name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, ascii_name: str, east_asia: str, size: float, color: str):
    style.font.name = ascii_name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_name)
    r_fonts.set(qn("w:hAnsi"), ascii_name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "CCD5DF", size: int = 6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_cols = list(table._tbl.tblGrid)
    for idx, width in enumerate(widths_dxa):
        if idx < len(grid_cols):
            grid_cols[idx].set(qn("w:w"), str(width))
        for row in table.rows:
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_keep(paragraph, *, with_next: bool = False, together: bool = False):
    paragraph.paragraph_format.keep_with_next = with_next
    paragraph.paragraph_format.keep_together = together
    paragraph.paragraph_format.widow_control = True


def add_numbering_definition(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    lvl.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), ASCII_FONT)
        r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
        r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_paragraph_keep(paragraph, together=True)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, ASCII_FONT, EAST_ASIA_FONT, 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    specs = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, MUTED, 0, 12),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        set_style_font(style, ASCII_FONT, EAST_ASIA_FONT, size, color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True


def configure_sections(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    set_header_footer(section)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "HORSE ASSET MANAGER  |  MVP DESIGN PACKAGE"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("INTERNAL DESIGN REFERENCE  •  ")
    set_run_font(run, size=8, color=MUTED)
    page_run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_char_begin)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char_end)
    set_run_font(page_run, size=8, color=MUTED)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    set_paragraph_keep(p, together=False)
    return p


def add_bullets(doc: Document, items: Iterable[str], bullet_num_id: int):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run)
        apply_numbering(p, bullet_num_id)


def add_numbers(doc: Document, items: Iterable[str], decimal_num_id: int):
    del decimal_num_id
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        run = p.add_run(f"{index}. {item}")
        set_run_font(run)


def add_callout(doc: Document, label: str, text: str, *, fill: str = CALLOUT, color: str = NAVY):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), "D5DEE8")
        borders.append(node)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, color=color, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    set_paragraph_keep(p, together=True)
    return p


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
    *,
    header_fill: str = LIGHT_BLUE,
    font_size: float = 9.3,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(header)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
        set_paragraph_keep(p, together=True)
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        if row_index % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "FAFBFC")
        for idx, value in enumerate(row_values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            if idx == 0 and len(headers) > 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK, bold=(idx == 0 and len(headers) > 1))
            set_paragraph_keep(p, together=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    set_paragraph_keep(p, together=True)
    return p


def add_figure(doc: Document, path: Path, caption: str, alt_text: str, width: float = 6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    set_paragraph_keep(p, together=True)
    inline_shape = doc.inline_shapes[-1]
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", caption)
    add_caption(doc, caption)


def page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def load_font(size: int, bold: bool = False):
    path = Path("C:/Windows/Fonts/YuGothB.ttc" if bold else "C:/Windows/Fonts/YuGothM.ttc")
    return ImageFont.truetype(str(path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    title: str,
    body: str,
    *,
    fill: str,
    outline: str = "#AAB8C7",
    title_color: str = "#0B2545",
):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=18, fill=fill, outline=outline, width=3)
    title_font = load_font(28, bold=True)
    body_font = load_font(21)
    title_lines = wrap_text(draw, title, title_font, x2 - x1 - 36)
    y = y1 + 22
    for line in title_lines:
        box = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((x1 + x2 - (box[2] - box[0])) / 2, y), line, font=title_font, fill=title_color)
        y += 36
    y += 8
    for line in wrap_text(draw, body, body_font, x2 - x1 - 40):
        box = draw.textbbox((0, 0), line, font=body_font)
        draw.text(((x1 + x2 - (box[2] - box[0])) / 2, y), line, font=body_font, fill="#3F4D5A")
        y += 30


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color="#4B6B88"):
    draw.line([start, end], fill=color, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = 0.55
    points = [
        end,
        (
            end[0] - length * math.cos(angle - spread),
            end[1] - length * math.sin(angle - spread),
        ),
        (
            end[0] - length * math.cos(angle + spread),
            end[1] - length * math.sin(angle + spread),
        ),
    ]
    draw.polygon(points, fill=color)


def make_architecture_diagram(path: Path):
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 35), "システム構成概要", font=load_font(38, bold=True), fill="#0B2545")
    draw_box(draw, (70, 160, 430, 430), "利用者ブラウザ", "React SPA / UI\nPDF.js端末内解析\nTanStack Query", fill="#F5F8FB")
    draw_box(draw, (620, 110, 1030, 350), "Cloudflare Worker", "Static Assets + Hono API\n認証・所有権・Zod検証", fill="#E8F0FA")
    draw_box(draw, (620, 500, 1030, 750), "日次メンテナンス", "JST 09:15\n予定補充・期限更新\n通知・セッション掃除", fill="#FFF6E3", outline="#D6B66C")
    draw_box(draw, (1210, 150, 1530, 390), "Cloudflare D1", "19テーブル\n利用者別データ分離\n円整数・監査", fill="#EDF7F1", outline="#8BB49A")
    draw_box(draw, (1210, 535, 1530, 735), "Workers Logs", "エラー・実行状況\n機密値は出力しない", fill="#F7F3FA", outline="#B6A1C6")
    arrow(draw, (430, 285), (620, 235))
    draw.text((470, 215), "HTTPS JSON", font=load_font(20, bold=True), fill="#4B6B88")
    arrow(draw, (1030, 235), (1210, 270))
    arrow(draw, (825, 500), (825, 350))
    arrow(draw, (1030, 625), (1210, 635))
    arrow(draw, (1030, 300), (1210, 600))
    draw.text((74, 775), "PDF本体・抽出全文はブラウザ外へ送信しない", font=load_font(25, bold=True), fill="#9B1C1C")
    draw.text((74, 822), "実績集計の正本は cashflows.status = confirmed", font=load_font(25, bold=True), fill="#2E6B4E")
    image.save(path, quality=95)


def make_dataflow_diagram(path: Path):
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 35), "資金データフロー", font=load_font(38, bold=True), fill="#0B2545")
    draw_box(draw, (70, 145, 380, 330), "入力", "手動登録\nPDF確認済み明細\n定期ルール", fill="#F5F8FB")
    draw_box(draw, (520, 95, 890, 300), "予定", "scheduled_cashflows\nplanned / paid\ncancelled / overdue", fill="#FFF6E3", outline="#D6B66C")
    draw_box(draw, (520, 485, 890, 690), "実績", "cashflows\nconfirmed / cancelled\narchived", fill="#EDF7F1", outline="#8BB49A")
    draw_box(draw, (1030, 300, 1320, 520), "照合", "1対1\n一致・差額\n未実績・予定外", fill="#E8F0FA")
    draw_box(draw, (1390, 145, 1560, 675), "出力", "ダッシュ\n台帳\n分析\nCSV\n予算", fill="#F7F3FA", outline="#B6A1C6")
    arrow(draw, (380, 225), (520, 195))
    arrow(draw, (380, 260), (520, 585))
    arrow(draw, (890, 205), (1030, 375))
    arrow(draw, (890, 585), (1030, 445))
    arrow(draw, (1320, 410), (1390, 410))
    arrow(draw, (890, 585), (1390, 565), color="#2E6B4E")
    draw.text((985, 620), "実績集計はconfirmedだけ", font=load_font(21, bold=True), fill="#2E6B4E")
    draw_box(draw, (70, 550, 380, 740), "契約・精算予定", "investments\nhorse_settlements\n実績とは別に保持", fill="#F2F4F7")
    arrow(draw, (380, 645), (520, 620))
    draw.text((72, 800), "予定・契約・精算予定を実績へ重ねて加算しない", font=load_font(27, bold=True), fill="#9B1C1C")
    image.save(path, quality=95)


def make_er_diagram(path: Path):
    image = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 30), "データモデル概要（19テーブル）", font=load_font(38, bold=True), fill="#0B2545")
    boxes = {
        "users": (70, 130, 340, 270, "利用者", "users / sessions"),
        "masters": (70, 390, 340, 590, "マスター", "clubs\ncategories\nbudgets"),
        "horses": (490, 100, 820, 310, "馬・出資", "horses\nhorse_name_aliases\ninvestments"),
        "cash": (490, 430, 820, 690, "資金", "cashflows\nrecurring_rules\nscheduled_cashflows\nreconciliations"),
        "planning": (970, 110, 1300, 330, "計画・精算", "simulation_scenarios\nsimulation_items\nhorse_settlements"),
        "ops": (970, 470, 1300, 720, "通知・監査", "alert_rules\nnotifications\naudit_logs\nstatement_imports"),
    }
    fills = ["#F5F8FB", "#F2F4F7", "#E8F0FA", "#EDF7F1", "#FFF6E3", "#F7F3FA"]
    for (key, (x1, y1, x2, y2, title, body)), fill in zip(boxes.items(), fills):
        draw_box(draw, (x1, y1, x2, y2), title, body, fill=fill)
    arrow(draw, (340, 200), (490, 205))
    arrow(draw, (340, 490), (490, 555))
    arrow(draw, (655, 310), (655, 430))
    arrow(draw, (820, 210), (970, 220))
    arrow(draw, (820, 555), (970, 580))
    arrow(draw, (1135, 330), (1135, 470))
    draw_box(draw, (70, 790, 1300, 920), "共通ルール", "主要表はuser_idで分離 / 金額は円整数 / 一意制約で冪等 / 馬以外は原則アーカイブ", fill="#FAFBFC")
    image.save(path, quality=95)


def generate_diagrams():
    ASSETS.mkdir(parents=True, exist_ok=True)
    architecture = ASSETS / "architecture-overview.png"
    dataflow = ASSETS / "financial-data-flow.png"
    er = ASSETS / "data-model-overview.png"
    make_architecture_diagram(architecture)
    make_dataflow_diagram(dataflow)
    make_er_diagram(er)
    return architecture, dataflow, er


def build_document():
    architecture_img, dataflow_img, er_img = generate_diagrams()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_id = add_numbering_definition(doc, bullet=True)
    decimal_id = add_numbering_definition(doc, bullet=False)

    props = doc.core_properties
    props.title = "Horse Asset Manager 設計資料一式"
    props.subject = "要件定義・構成図・DB/API/UI・セキュリティ・テスト・運用"
    props.author = "Horse Asset Manager Project"
    props.keywords = "一口馬主, 資金管理, Cloudflare, D1, React, 設計"
    props.comments = "Generated from repository design sources."

    # Cover: editorial_cover pattern with a compact Japanese font override.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(80)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("MVP DESIGN PACKAGE")
    set_run_font(run, size=10.5, color=GOLD, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Horse Asset Manager")
    title2 = doc.add_paragraph(style="Title")
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.paragraph_format.space_after = Pt(10)
    title2.add_run("設計資料一式")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("要件定義・構成図・DB/API/UI・セキュリティ・テスト・運用")
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(74)
    r = line.add_run("一口馬主活動の、過去と未来のお金を管理する")
    set_run_font(r, size=11, color=GOLD, bold=True)
    for text in (
        "文書版 1.1",
        "基準日 2026年8月12日",
        "対象実装 自分専用安定運用版（2026-08-12作業ツリー）",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        rr = p.add_run(text)
        set_run_font(rr, size=10, color=MUTED, bold=text.startswith("文書版"))
    doc.add_paragraph().paragraph_format.space_after = Pt(28)
    add_callout(
        doc,
        "スコープ",
        "本アプリは競馬予想・投資推奨ではありません。出資金、維持費、会費、保険料、分配金、精算金と将来予定を管理します。",
        fill="FFF8E8",
        color=GOLD,
    )
    page_break(doc)

    doc.add_heading("文書ガイド", level=1)
    add_body(
        doc,
        "本Word版はレビュー・配布用の統合概要です。詳細と更新の正本はリポジトリの docs/README.md 以下にあるMarkdown資料です。",
    )
    add_table(
        doc,
        ["項目", "内容"],
        [
            ["対象", "Windows Web版MVP / Chrome・Edge"],
            ["現状", "本人限定の安定運用改善まで実装済み。一般公開prodは延期"],
            ["技術", "React + TypeScript + Vite / Hono Worker / D1 / Drizzle / Zod"],
            ["最重要ルール", "実績集計はcashflowsのconfirmedのみ。金額は円整数"],
            ["個人運用判定", "Access本人限定、登録停止、D1統合、バックアップ・復元を必須"],
        ],
        [1900, 7460],
    )
    doc.add_heading("収録内容", level=2)
    add_numbers(
        doc,
        [
            "プロジェクト概要と業務ルール",
            "MVP要件と受入条件",
            "システム構成・データフロー",
            "データベース・API・画面設計",
            "セキュリティ・データ保護",
            "テスト・運用・コスト設計",
            "ロードマップ・主要リスク・未決事項",
        ],
        decimal_id,
    )
    add_callout(doc, "安定運用の重点", "精算完了は状態条件・利用者単位の冪等性キー・D1 batchで保護し、連続・同時再実行を409で拒否します。一般公開向け認証・規約・prod構築は延期します。", fill="EDF7F1", color=GREEN)

    doc.add_heading("1. プロジェクト概要", level=1)
    add_body(
        doc,
        "Horse Asset Managerは、一口馬主活動における資金移動を、出資検討から引退・精算完了まで一元管理するWebサービスです。過去の記録だけでなく、支払い予定、予算超過、出資可能額を見える化します。",
    )
    doc.add_heading("1.1 利用者価値", level=2)
    add_bullets(
        doc,
        [
            "今月と今後12か月の支出予定を把握できる。",
            "年間・月間予算に対する実績、未払い予定、見込、残額を分けて確認できる。",
            "候補馬の初回、月額、初年度、想定期間の資金影響を比較できる。",
            "馬・クラブ・カテゴリー・月ごとの支出、入金、損益、回収率を確認できる。",
            "馬名変更や引退後も同じ馬IDで精算完了まで追跡できる。",
        ],
        bullet_id,
    )
    doc.add_heading("1.2 スコープ", level=2)
    add_table(
        doc,
        ["MVP対象", "MVP対象外"],
        [
            ["Windows 10/11、Chrome、Edge", "iOS/Android/Swiftネイティブ"],
            ["出資・収支・予定・予算・精算", "レース・馬券予想"],
            ["馬別・クラブ別・カテゴリー別分析", "血統・調教・馬体評価"],
            ["対応PDFのブラウザ内解析", "AI OCR・画像/暗号化/未対応PDF"],
            ["アプリ内通知、CSV", "課金、一般公開、コミュニティ"],
        ],
        [4680, 4680],
    )
    doc.add_heading("1.3 重要な業務ルール", level=2)
    add_table(
        doc,
        ["ID", "ルール"],
        [
            ["BR-01", "金額は円単位の安全な整数。小数を保存しない。"],
            ["BR-02", "実績集計の正本はcashflows.status=confirmedだけ。"],
            ["BR-03", "馬代回収率=入金÷出資元本、総合回収率=入金÷実績支出。分母0は算出不可。"],
            ["BR-04", "出資可能額=max(0, 年間予算-実績支出-未払い予定支出)。"],
            ["BR-05", "馬ステータスは整理ラベル。状態変更だけで金融データを変えない。"],
            ["BR-06", "現在名と旧名を保持し、収支は馬IDで結び付ける。"],
            ["BR-07", "馬だけは名前完全一致後に関連データを完全削除。通常はアーカイブ。"],
            ["BR-08", "PDF本体・抽出全文はサーバーへ送信・保存しない。"],
        ],
        [1350, 8010],
        font_size=9.0,
    )

    doc.add_heading("2. MVP要件", level=1)
    add_body(doc, "要件IDは設計、実装、テスト、課題管理で共通利用します。詳細な受入条件は docs/01_requirements.md を正本とします。")
    add_table(
        doc,
        ["領域", "主要要件", "状態"],
        [
            ["認証・初期設定", "登録許可設定、ログイン、Cookieセッション、初期予算・マスター", "実装済み"],
            ["馬・出資", "候補/出資馬、旧名、出資条件、初回支出、完全削除", "実装済み"],
            ["実績収支", "支出・入金、絞込、編集、アーカイブ", "実装済み"],
            ["予定・照合", "定期ルール、12か月予定、候補比較、1対1照合・解除", "実装済み"],
            ["予算・計画", "年/月予算、出資可能額、複数候補シミュレーション", "実装済み"],
            ["分析・台帳", "ダッシュボード、各軸分析、回収率、安全化CSV", "実装・D1一致試験済み"],
            ["引退・精算", "精算予定、冪等な実績反映、馬の精算完了", "実装・並行試験済み"],
            ["通知", "期限、締切、予算、未入力、集中、既読", "実装済み"],
            ["PDF", "silk/lord対応、端末内解析、確認、一括登録、重複拒否", "実装済み"],
        ],
        [1900, 5860, 1600],
        font_size=8.8,
    )
    doc.add_heading("2.1 主要受入シナリオ", level=2)
    add_numbers(
        doc,
        [
            "新規登録→初期設定→候補馬→出資＋初回支出→馬別台帳まで完了できる。",
            "定期ルールから予定を生成し、再実行で重複せず、実績と照合できる。",
            "対応PDFを確認して予定/実績へ登録し、同じPDFの再取込を拒否する。",
            "別利用者の既知IDを使っても詳細・更新・削除・分析・CSVへ混入しない。",
            "馬名完全一致時だけ関連データを完全削除し、匿名件数監査だけを残す。",
            "精算を一度だけ実績へ反映し、未完了がない場合だけ馬を精算完了にする。",
        ],
        decimal_id,
    )

    doc.add_heading("3. システム構成", level=1)
    add_figure(
        doc,
        architecture_img,
        "図1. システム構成概要",
        "利用者ブラウザ、Cloudflare Worker、日次処理、D1、Workers Logsの関係。PDFはブラウザ内で解析する。",
    )
    add_table(
        doc,
        ["レイヤー", "責務", "配置"],
        [
            ["Web", "画面、フォーム、ルーティング、チャート、PDF解析", "apps/web"],
            ["API", "認証、Origin、所有権、Zod、集計、CSV", "apps/api"],
            ["共有", "金額、回収率、予算、日付、予定生成", "packages/shared"],
            ["入力契約", "Web/API共通Zod", "packages/validation"],
            ["DB", "Drizzle schema、D1接続、制約・索引", "packages/database"],
            ["運用", "日次予定補充、期限、通知、セッション掃除", "Worker Cron"],
        ],
        [1800, 5000, 2560],
        font_size=9.0,
    )
    doc.add_heading("3.1 資金データフロー", level=2)
    add_figure(
        doc,
        dataflow_img,
        "図2. 予定・実績・契約・精算の分離",
        "入力から予定、実績、照合、分析への流れ。実績集計はconfirmed cashflowsだけを使う。",
    )

    doc.add_heading("4. データベース設計", level=1)
    add_figure(
        doc,
        er_img,
        "図3. データモデル概要",
        "19テーブルを利用者、マスター、馬・出資、資金、計画・精算、通知・監査の領域へ分類した概要。",
    )
    add_table(
        doc,
        ["領域", "テーブル", "重要制約"],
        [
            ["利用者", "users, sessions", "メール一意、セッション期限索引"],
            ["マスター", "clubs, categories, budgets", "利用者内名称・期間一意、金額非負"],
            ["馬・出資", "horses, horse_name_aliases, investments", "旧名一意、1馬1出資、口数正数"],
            ["取込・実績", "statement_imports, cashflows", "PDF hash一意、取込行一意、金額非負"],
            ["予定・照合", "recurring_rules, scheduled_cashflows, cashflow_reconciliations", "ルール+期日一意、照合1対1"],
            ["計画・精算", "simulation_scenarios, simulation_items, horse_settlements", "明細金額非負、cashflow link一意"],
            ["通知・監査", "alert_rules, notifications, audit_logs", "利用者+種別/通知key一意、機密値をredact"],
        ],
        [1500, 3920, 3940],
        font_size=8.8,
    )
    doc.add_heading("4.1 完全削除", level=2)
    add_body(doc, "馬の完全削除は、現在名の完全一致確認後、D1 batchで次の依存順に行います。")
    add_numbers(
        doc,
        [
            "馬名・馬ID・金額を含まない匿名のテーブル別件数監査を先に作成する。",
            "対象通知、馬の詳細監査、予定・実績照合、精算を削除する。",
            "予定、定期ルール、シミュレーション明細、実績、出資、旧名を削除する。",
            "最後に馬本体を削除する。シミュレーションのシナリオ本体は残す。",
        ],
        decimal_id,
    )
    add_callout(doc, "注意", "復元機能はありません。運用者が利用者の代わりに直接SQLで馬を削除しない方針です。", fill="FFF8E8", color=GOLD)

    doc.add_heading("5. API設計", level=1)
    add_body(doc, "APIは /api 配下のREST形式です。成功は data/message、一覧はmeta、エラーはcode/message/detailsで返します。入力不正422、業務競合409、他利用者を含む未検出404を基本とします。")
    add_table(
        doc,
        ["API群", "代表パス", "設計ポイント"],
        [
            ["認証", "/auth/config, /register, /login, /logout, /me", "登録許可設定、14日Cookie"],
            ["初期設定", "/setup, /setup/defaults", "予算・マスター・アラート一括"],
            ["馬・出資", "/horses, /investments", "旧名、所有権、完全削除"],
            ["収支", "/cashflows", "実績正本、最大100件ページング"],
            ["予定", "/recurring-rules, /scheduled-cashflows", "12か月、ON CONFLICTで冪等"],
            ["照合", "/reconciliations, /auto-match", "予定/実績の1対1、解除・状態復帰"],
            ["分析", "/dashboard, /analytics, /ledger", "期間限定、confirmed集計"],
            ["精算", "/settlements/:id/complete", "batch・冪等性キー、再実行409"],
            ["PDF", "/statement-imports", "最大100明細、hash/行key一意"],
            ["出力", "/export/*.csv", "最大5年、UTF-8 BOM、式注入対策"],
        ],
        [1700, 3330, 4330],
        font_size=8.7,
    )
    doc.add_heading("5.1 API共通保護", level=2)
    add_bullets(
        doc,
        [
            "Hono secureHeadersを全ルートへ適用する。",
            "GET/HEAD/OPTIONS以外はOriginを検証する。",
            "認証利用者はCookieから決定し、本文のuser_idを信用しない。",
            "馬・クラブ・カテゴリーなどの参照先を同じuser_idで確認する。",
            "一覧は原則ページングし、pageSizeは最大100とする。",
        ],
        bullet_id,
    )

    doc.add_heading("6. UI・画面設計", level=1)
    add_body(doc, "PCでは248pxの左サイドバー、1024px未満では上部ヘッダーと開閉メニューを使います。360px以上で主要操作を可能にし、狭い画面では表をカードへ切り替えます。")
    add_table(
        doc,
        ["領域", "主なルート", "主要操作"],
        [
            ["認証", "/login, /register, /setup", "登録、ログイン、初期予算"],
            ["概要", "/dashboard, /notifications", "月次・年次見込、予定、通知"],
            ["馬", "/prospects, /horses, /horses/:id", "候補、出資、台帳、精算、削除"],
            ["収支", "/cashflows, /cashflows/import", "実績、PDF確認・一括登録"],
            ["予定", "/scheduled, /calendar, /reconciliations", "ルール、予定、照合"],
            ["計画", "/budgets, /simulations", "予算、出資可能額、候補比較"],
            ["分析", "/analytics", "期間と軸、損益・回収率"],
            ["設定", "/settings/*", "クラブ、カテゴリー、アラート、CSV"],
        ],
        [1500, 3750, 4110],
        font_size=8.9,
    )
    doc.add_heading("6.1 UI受入", level=2)
    add_bullets(
        doc,
        [
            "予定・実績・見込を文言で区別し、色だけに頼らない。",
            "入力エラーを項目近くに日本語で示し、APIエラーは内部情報を見せない。",
            "馬の完全削除は影響を表示し、名前一致まで実行ボタンを無効にする。",
            "キーボード、フォーカス、ラベル、ダイアログ、200%ズームを確認する。",
            "WindowsのChrome・Edgeで主要6シナリオを完了する。",
        ],
        bullet_id,
    )

    doc.add_heading("7. セキュリティ・データ保護", level=1)
    add_table(
        doc,
        ["領域", "個人運用で実装済み", "一般公開時・継続対応"],
        [
            ["認証", "PBKDF2、hash session、Cookie、Access、登録停止", "メール確認、reset、rate limit、IdP再評価"],
            ["認可", "全主要表user_id、他者404、主要API D1マトリクス", "新規API追加時の回帰"],
            ["Web", "secureHeaders、更新Origin確認、Access(dev)", "CSP、欠落Origin/CSRF方針、HSTS"],
            ["入力", "strict Zod、prepared SQL、DB制約、失敗注入D1統合", "制約境界と大量データを継続"],
            ["PDF", "端末内解析、API最小payload、hash重複拒否", "Network/D1/ログ回帰E2E"],
            ["ログ", "requestId・処理時間・エラー種別だけ", "実環境tail、保持、定期監査"],
            ["CSV", "最大5年、BOM、引用、式注入対策とD1試験", "Excel更新時の回帰"],
        ],
        [1450, 3955, 3955],
        font_size=8.5,
    )
    add_callout(doc, "最高機密", "PDF本体・抽出全文、パスワード、セッションCookieは、ログや監査JSONへ残しません。", fill="FCEBEC", color=RED)

    doc.add_heading("8. テスト・受入", level=1)
    add_body(doc, "単体→API→D1統合→Playwright E2E→手動受入の順で、金額、権限、冪等性、破壊的操作、ブラウザを縦断確認します。")
    add_table(
        doc,
        ["シナリオ", "確認事項", "現状"],
        [
            ["初回利用", "登録→設定→候補→出資→収支→台帳", "E2E済み"],
            ["利用者分離", "別利用者の主要リソースを404", "D1所有権マトリクス済み"],
            ["予定・照合", "月末、12か月、再生成、差額、解除", "D1・E2E済み"],
            ["PDF", "合計、紐付け、hash、非送信", "parser単体 / Network E2E追加"],
            ["精算", "実績反映は1回、pending中は馬完了不可", "連続・同時D1/E2E済み"],
            ["完全削除", "不一致拒否、依存削除、匿名監査", "候補/出資馬E2E済み"],
            ["大量・CSV", "BOM・式注入・集計一致", "安全化済み / 10,000件は継続"],
        ],
        [1800, 5100, 2460],
        font_size=8.7,
    )
    doc.add_heading("8.1 品質ゲート", level=2)
    add_bullets(
        doc,
        [
            "typecheck、lint、format:check、test、D1統合、build、Chrome/Edge E2Eが成功する。",
            "19テーブル、全migration、foreign_key_check、unique/check、batch失敗を確認する。",
            "固定金額データでdashboard、ledger、analytics、CSVが一致する。",
            "P0/P1不具合が0件である。",
            "Access・登録停止・ログ・SQLバックアップを確認し、次回DB変更前にTime Travelを訓練する。",
        ],
        bullet_id,
    )

    doc.add_heading("9. 運用・コスト", level=1)
    add_table(
        doc,
        ["環境", "Worker / D1", "用途"],
        [
            ["local", "local Worker / local D1", "開発・自動テスト"],
            ["dev", "horse-asset-manager-dev / horse_asset_manager_dev", "Access限定の受入・運用確認"],
            ["prod", "未構築・devと完全分離", "一般公開を決めた場合だけ"],
        ],
        [1500, 4400, 3460],
        font_size=9.0,
    )
    doc.add_heading("9.1 リリースと復旧", level=2)
    add_numbers(
        doc,
        [
            "コード品質とE2Eを完了し、D1復旧基準点を記録する。",
            "devへmigrationを適用し、foreign_key_checkとdry-runを行う。",
            "deploy後にhealth、Access、認証、主要画面、Cookie、深いSPAルートを確認する。",
            "24時間、5xx、Cron、D1 read/write、CPU、費用を監視する。",
            "障害時はWorker版とDB互換を確認してrollbackし、必要ならTime Travelで復旧する。",
        ],
        decimal_id,
    )
    doc.add_heading("9.2 Cloudflare費用前提", level=2)
    add_table(
        doc,
        ["項目", "2026-08-12時点の前提", "ガードレール"],
        [
            ["Workers", "Paid最低5 USD/月 + 超過従量", "CPU上限、再試行制御、Budget Alert"],
            ["D1 read", "月250億行含有", "期間、pageSize<=100、索引、rows_read監視"],
            ["D1 write", "月5,000万行含有", "一意制約、dedupe、日次Cron"],
            ["D1 storage", "月5GB含有、1 DB上限10GB", "容量月次確認、巨大JSON/PDFを保存しない"],
            ["復旧", "日時付きSQL + Paid Time Travel", "local復元済み、dev訓練は次回DB変更前"],
        ],
        [1700, 3400, 4260],
        font_size=8.6,
    )
    add_body(doc, "公式情報: developers.cloudflare.com/workers/platform/pricing/ および developers.cloudflare.com/d1/platform/pricing/・limits/。公開前に再確認します。")

    doc.add_heading("10. ロードマップ・主要リスク", level=1)
    add_table(
        doc,
        ["状態", "作業", "内容"],
        [
            ["完了", "精算完了の再実行防止", "状態条件・冪等性キー・再送409・D1並行試験"],
            ["完了", "D1 batch統合", "出資/PDF/定期予定/精算の失敗注入"],
            ["完了", "訂正UIとCSV", "照合解除・精算状態・BOM・式注入対策"],
            ["完了", "個人運用保護", "Access本人限定、登録停止、最小ログ"],
            ["一部", "復旧訓練", "local SQL復元済み。dev Time Travelは次回変更前"],
            ["継続", "A11y・大量データ", "手動監査、10,000収支、5年性能"],
            ["延期", "一般公開対応", "メール確認、reset、rate limit、規約、prod分離"],
        ],
        [1050, 3260, 5050],
        font_size=8.7,
    )
    doc.add_heading("10.1 採用した主要判断", level=2)
    add_bullets(
        doc,
        [
            "資金管理に限定し、競馬予想・投資推奨を実装しない。",
            "React SPAとHono APIを同じWorkerから配信する。",
            "環境ごとに1つのD1を使い、行のuser_idで分離する。",
            "予定・実績・出資契約・精算予定を分離する。",
            "馬ステータスは自由ラベル、馬名は旧名履歴を保持する。",
            "馬だけは明示確認後に完全削除する。",
            "PDFは端末内解析し、本体と全文を送らない。",
        ],
        bullet_id,
    )

    doc.add_heading("付録A. 詳細資料", level=1)
    add_body(doc, "次のMarkdown資料が詳細と更新の正本です。")
    add_table(
        doc,
        ["ファイル", "内容"],
        [
            ["docs/README.md", "資料索引・変更管理"],
            ["docs/00_project_overview.md", "目的・スコープ・業務ルール"],
            ["docs/01_requirements.md", "機能・非機能要件と受入"],
            ["docs/02_architecture.md", "構成・配置・シーケンス・境界"],
            ["docs/03_database_design.md", "19テーブル・制約・削除"],
            ["docs/04_api_design.md", "全API・認証・エラー・冪等"],
            ["docs/05_ui_routes.md", "ルート・主要画面・A11y"],
            ["docs/06_cost_and_cloudflare_paid_plan.md", "価格・制限・監視"],
            ["docs/07_implementation_plan.md", "現在地・P0/P1・品質ゲート"],
            ["docs/08_security_and_data_protection.md", "脅威・保護・残対応"],
            ["docs/09_test_and_acceptance_plan.md", "テスト戦略・受入"],
            ["docs/10_operations_and_release.md", "デプロイ・障害・復旧"],
            ["docs/11_requirements_traceability.md", "要件対応表"],
            ["docs/12_decisions_and_open_issues.md", "ADR・リスク・未決事項"],
        ],
        [4100, 5260],
        font_size=8.6,
    )
    doc.add_heading("付録B. 公式参照", level=1)
    add_bullets(
        doc,
        [
            "Workers Pricing — https://developers.cloudflare.com/workers/platform/pricing/",
            "D1 Pricing — https://developers.cloudflare.com/d1/platform/pricing/",
            "D1 Limits — https://developers.cloudflare.com/d1/platform/limits/",
            "D1 Time Travel — https://developers.cloudflare.com/d1/reference/time-travel/",
            "Workers Logs — https://developers.cloudflare.com/workers/observability/logs/workers-logs/",
        ],
        bullet_id,
    )
    add_callout(doc, "文書管理", "要件、API、DB、画面、テストの変更時は、同じ変更単位で詳細資料とトレーサビリティを更新します。", fill="EAF3FB", color=BLUE)

    # Ensure all inline runs in title styles carry Japanese font override.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                set_run_font(run)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
